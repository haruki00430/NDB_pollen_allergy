"""Fix What is new? bullet order and Study Registration heading style in JCE DOCX."""
from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument

DOCX_PATH = Path(__file__).resolve().parent / (
    "How_Outcome_Definition_Changes_Ecological_Inference_20260618.docx"
)

WHAT_IS_NEW_BULLETS = [
    (
        "Outcome definition changed recoverable ecological signal by 1.72-fold "
        "(R\u00b2 = 0.469 vs. 0.272) under identical exposure data, covariates, "
        "and regression frameworks."
    ),
    (
        "Pollen associations remained significant in 7/8 sensitivity "
        "specifications for the allergy-specific outcome but only 5/6 for the "
        "best comparator model, demonstrating specification-dependent attenuation "
        "with heterogeneous outcomes."
    ),
    (
        "A pharmacologically unrelated negative control (oral hypoglycaemic "
        "agents) showed no pollen association in any model, confirming that "
        "signal differences reflected outcome specificity rather than generic "
        "ecological confounding."
    ),
    (
        "Medication dispensing counts function as measurement instruments "
        "rather than interchangeable disease proxies in aggregate prescription "
        "data."
    ),
    (
        "Researchers using claims or open prescription aggregates should evaluate "
        "pharmacological specificity before model selection, interaction testing, "
        "or spatial refinement."
    ),
]


def find_heading(doc: DocxDocument, text: str, level: int = 1) -> Paragraph:
    style = f"Heading {level}"
    for para in doc.paragraphs:
        style_name = para.style.name if para.style is not None else None
        if style_name == style and para.text.strip() == text:
            return para
    raise RuntimeError(f"Heading not found: {text!r}")


def remove_paragraph(para: Paragraph) -> None:
    element = para._element
    element.getparent().remove(element)


def main() -> None:
    doc = Document(str(DOCX_PATH))

    heading = find_heading(doc, "What is new?")
    intro = find_heading(doc, "Introduction")
    heading_index = None
    intro_index = None
    for idx, para in enumerate(doc.paragraphs):
        if para._p is heading._p:
            heading_index = idx
        if para._p is intro._p:
            intro_index = idx
    if heading_index is None or intro_index is None:
        raise RuntimeError("Could not locate section boundaries")

    for para in list(doc.paragraphs[heading_index + 1 : intro_index]):
        remove_paragraph(para)

    intro = find_heading(doc, "Introduction")
    new_paras = []
    for bullet in WHAT_IS_NEW_BULLETS:
        para = doc.add_paragraph()
        para.add_run("\u2022 " + bullet)
        new_paras.append(para)

    for para in new_paras:
        intro._p.addprevious(para._p)

    for para in doc.paragraphs:
        style_name = para.style.name if para.style is not None else ""
        if para.text.strip() == "Study Registration" and style_name != "Heading 2":
            para.style = doc.styles["Heading 2"]  # type: ignore[assignment]
            break

    doc.save(str(DOCX_PATH))
    print(f"Fixed: {DOCX_PATH}")


if __name__ == "__main__":
    main()
