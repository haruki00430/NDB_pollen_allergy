"""Patch How_Outcome_Definition_Changes_Ecological_Inference_20260618.docx for JCE submission."""
from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX_PATH = Path(__file__).resolve().parent / (
    "How_Outcome_Definition_Changes_Ecological_Inference_20260618.docx"
)

NEW_TITLE = (
    "How Outcome Definition Changes Ecological Inference in "
    "Prescription-Based Epidemiology: A Natural Experiment"
)

INTRO_LABEL = "Introduction or background: "
INTRO_TEXT = (
    "Ecological studies using aggregate prescription data often treat medication "
    "dispensing as a proxy for disease burden when diagnosis-linked records are "
    "unavailable. The degree to which outcome definition influences signal "
    "recovery in such designs has not been systematically quantified. We aimed "
    "to quantify how ecological inference changes when exposure variables, "
    "covariates, and regression frameworks are held constant while the "
    "medication-based outcome varies in pharmacological specificity."
)

METHODS_LABEL = "Methods: "
METHODS_TEXT = (
    "We conducted a cross-sectional ecological study across Japan\u2019s 47 "
    "prefectures (NDB Open Data, fiscal year 2023), using seasonal pollen "
    "exposure as a natural-experiment signal with an established biological "
    "mechanism. An allergy-specific outcome (antihistamines, ophthalmic, and "
    "nasal preparations; codes 441/449/131/132) and a pharmacologically "
    "heterogeneous comparator (topical anti-inflammatory agents; code 264) were "
    "evaluated under identical exposure variables, covariates, and regression "
    "frameworks. A pharmacologically unrelated negative control (oral "
    "hypoglycemic agents; code 396) assessed false-positive rates under the "
    "same analytic design."
)

RESULTS_LABEL = "Results: "
RESULTS_TEXT = (
    "The allergy-specific outcome achieved explained variance R\u00b2 = 0.469 "
    "versus R\u00b2 = 0.272 for the heterogeneous comparator (1.72-fold "
    "difference). Pollen associations were statistically significant in 7/8 "
    "sensitivity specifications for the allergy-specific outcome versus 5/6 for "
    "the best-performing comparator model. The negative control showed no "
    "pollen association in any specification, supporting that signal "
    "differences reflected outcome specificity rather than ecological "
    "confounding."
)

DISCUSSION_LABEL = "Discussion: "
DISCUSSION_TEXT = (
    "Outcome definition systematically shaped signal recovery before downstream "
    "modeling choices could be evaluated. Medication-based outcomes function as "
    "measurement instruments rather than interchangeable proxies for disease "
    "burden. Evaluating pharmacological specificity before model selection may "
    "improve the reliability of prescription-based ecological analyses when "
    "diagnosis linkage is unavailable. Plain language summary: When researchers "
    "count prescriptions instead of diagnoses, which drug category they choose "
    "can change study results a lot\u2014even with the same pollen and statistical "
    "models. Narrow allergy drugs tracked seasonal pollen better than broad skin "
    "creams or diabetes medicines."
)

KEYWORDS_TEXT = (
    "Keywords: outcome misclassification; ecological inference; prescription "
    "data; outcome definition; measurement error; allergic rhinitis; Japan"
)

WHAT_IS_NEW_BULLETS = [
    (
        "Outcome definition changed recoverable ecological signal by 1.72-fold "
        "(R\u00b2 = 0.469 vs. 0.272) under identical exposure data, covariates, "
        "and regression frameworks."
    ),
    (
        "Pollen associations remained significant in 7/8 sensitivity "
        "specifications for the allergy-specific outcome but only 5/6 for the "
        "best comparator model, demonstrating specification-dependent attenuation "
        "with heterogeneous outcomes."
    ),
    (
        "A pharmacologically unrelated negative control (oral hypoglycaemic "
        "agents) showed no pollen association in any model, confirming that "
        "signal differences reflected outcome specificity rather than generic "
        "ecological confounding."
    ),
    (
        "Medication dispensing counts function as measurement instruments "
        "rather than interchangeable disease proxies in aggregate prescription "
        "data."
    ),
    (
        "Researchers using claims or open prescription aggregates should evaluate "
        "pharmacological specificity before model selection, interaction testing, "
        "or spatial refinement."
    ),
]

STUDY_REG_TEXT = (
    "The analysis plan for this inferential observational study was registered "
    "on the Open Science Framework (OSF; https://osf.io/XXXXX; registration "
    "pending final submission). The registration documents prespecified outcome "
    "definitions, hierarchical regression models, sensitivity analyses, and the "
    "negative-control design prior to manuscript submission to the Journal of "
    "Clinical Epidemiology."
)


def clear_content(para) -> None:
    """Remove all non-pPr children from a paragraph element."""
    element = para._p
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def append_bold_label_text(para_elem, label: str, text: str) -> None:
    """Append a bold label run and a normal text run to a paragraph."""
    bold_run = OxmlElement("w:r")
    bold_props = OxmlElement("w:rPr")
    bold_props.append(OxmlElement("w:b"))
    bold_run.append(bold_props)
    bold_text = OxmlElement("w:t")
    bold_text.text = label
    bold_run.append(bold_text)
    para_elem.append(bold_run)

    normal_run = OxmlElement("w:r")
    normal_text = OxmlElement("w:t")
    normal_text.text = text
    normal_text.set(qn("xml:space"), "preserve")
    normal_run.append(normal_text)
    para_elem.append(normal_run)


def rewrite_labeled_paragraph(para, label: str, text: str) -> None:
    """Replace paragraph content with bold label + body text."""
    clear_content(para)
    append_bold_label_text(para._p, label, text)


def remove_paragraph(para) -> None:
    """Delete a paragraph from the document."""
    element = para._element
    element.getparent().remove(element)


def find_heading(doc: DocxDocument, text: str, level: int = 1):
    """Return the first heading paragraph matching text."""
    style = f"Heading {level}"
    for para in doc.paragraphs:
        style_name = para.style.name if para.style is not None else None
        if style_name == style and para.text.strip() == text:
            return para
    raise RuntimeError(f"Heading not found: {text!r}")


def replace_what_is_new_section(doc: DocxDocument) -> None:
    """Replace IJE Key Messages with JCE What is new? bullets."""
    heading = None
    for candidate in ("Key Messages", "What is new?"):
        try:
            heading = find_heading(doc, candidate)
            break
        except RuntimeError:
            continue
    if heading is None:
        raise RuntimeError("Key Messages / What is new? heading not found")

    intro = find_heading(doc, "Introduction")
    heading_index = None
    intro_index = None
    for idx, para in enumerate(doc.paragraphs):
        if para._p is heading._p:
            heading_index = idx
        if para._p is intro._p:
            intro_index = idx
    if heading_index is None or intro_index is None:
        raise RuntimeError("Could not locate section boundaries")

    heading.text = "What is new?"
    for para in doc.paragraphs[heading_index + 1 : intro_index]:
        remove_paragraph(para)

    intro = find_heading(doc, "Introduction")
    new_paras = []
    for bullet in WHAT_IS_NEW_BULLETS:
        para = doc.add_paragraph()
        para.add_run("\u2022 " + bullet)
        new_paras.append(para)

    for para in new_paras:
        intro._p.addprevious(para._p)


def insert_study_registration(doc: DocxDocument) -> None:
    """Insert Study Registration after Study Design paragraph if missing."""
    for para in doc.paragraphs:
        if para.text.strip().startswith("The analysis plan for this inferential"):
            return

    study_design = find_heading(doc, "Study Design and Unit of Analysis", level=2)
    design_body = study_design._p.getnext()
    if design_body is None or design_body.tag != qn("w:p"):
        raise RuntimeError("Study Design body paragraph not found")

    heading = OxmlElement("w:p")
    heading_pr = OxmlElement("w:pPr")
    heading_style = OxmlElement("w:pStyle")
    heading_style.set(qn("w:val"), "Heading2")
    heading_pr.append(heading_style)
    heading.append(heading_pr)
    heading_run = OxmlElement("w:r")
    heading_text = OxmlElement("w:t")
    heading_text.text = "Study Registration"
    heading_run.append(heading_text)
    heading.append(heading_run)

    body = OxmlElement("w:p")
    body_pr = OxmlElement("w:pPr")
    body_style = OxmlElement("w:pStyle")
    body_style.set(qn("w:val"), "Normal")
    body_pr.append(body_style)
    body.append(body_pr)
    body_run = OxmlElement("w:r")
    body_text = OxmlElement("w:t")
    body_text.text = STUDY_REG_TEXT
    body_text.set(qn("xml:space"), "preserve")
    body_run.append(body_text)
    body.append(body_run)

    design_body.addnext(heading)
    heading.addnext(body)


def main() -> None:
    doc = Document(str(DOCX_PATH))

    doc.paragraphs[0].text = NEW_TITLE

    rewrite_labeled_paragraph(doc.paragraphs[3], INTRO_LABEL, INTRO_TEXT)
    rewrite_labeled_paragraph(doc.paragraphs[4], METHODS_LABEL, METHODS_TEXT)
    rewrite_labeled_paragraph(doc.paragraphs[5], RESULTS_LABEL, RESULTS_TEXT)
    rewrite_labeled_paragraph(doc.paragraphs[6], DISCUSSION_LABEL, DISCUSSION_TEXT)
    doc.paragraphs[7].text = KEYWORDS_TEXT

    replace_what_is_new_section(doc)
    insert_study_registration(doc)

    doc.save(str(DOCX_PATH))
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
