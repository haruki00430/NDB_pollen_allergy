"""
How_Outcome_Definition_Changes_Ecological_Inference_20260617.docx の段落構造を診断する。
PDS用DOCX修正のための参照スクリプト。
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

SRC = (
    r"C:\Users\user\.ag-cursor-common\research_workspace\projects"
    r"\NDB_Research_Hub\projects\NDB_XXX_pollen_allergy_v2\04_Manuscripts"
    r"\How_Outcome_Definition_Changes_Ecological_Inference_20260617.docx"
)

doc = Document(SRC)
print(f"総段落数: {len(doc.paragraphs)}")
print(f"総テーブル数: {len(doc.tables)}")
print()
print("=== 全段落（空白除く） ===")
for i, para in enumerate(doc.paragraphs):
    t = ''.join(r.text for r in para.runs)
    if t.strip():
        style = para.style.name if para.style else "Normal"
        print(f"Para#{i:3d} [{style:20s}]: {t[:120]!r}")
