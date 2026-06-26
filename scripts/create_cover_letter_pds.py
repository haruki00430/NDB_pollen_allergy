"""
PDS用カバーレター DOCX を作成する。
submission_package_PDS/CoverLetter_PDS.docx として保存。
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DST = (
    r"C:\Users\user\.ag-cursor-common\research_workspace\projects"
    r"\NDB_Research_Hub\projects\NDB_XXX_pollen_allergy_v2\04_Manuscripts"
    r"\submission_package_PDS\CoverLetter_PDS.docx"
)

doc = Document()

# ページ余白の設定
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def add_para(text, bold=False, font_size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    return p

def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p

# ── 日付・差出人情報 ──────────────────────────────────────────────────────────
add_para("26 June 2026", space_after=12)

add_para("The Editor-in-Chief", space_after=0)
add_para("Pharmacoepidemiology and Drug Safety", bold=True, space_after=0)
add_para("Wiley", space_after=12)

# ── 件名 ─────────────────────────────────────────────────────────────────────
add_para("Re: Submission of Original Article", bold=True, space_after=2)
add_para(
    '"How Outcome Definition Changes Ecological Pharmacoepidemiological Inference: '
    'A Natural Experiment on Outcome Misclassification"',
    font_size=11, space_after=12
)

# ── 本文 ─────────────────────────────────────────────────────────────────────
add_para("Dear Editor-in-Chief,", space_after=6)

add_para(
    "I am pleased to submit the above-titled manuscript for consideration for publication "
    "in Pharmacoepidemiology and Drug Safety as an Original Article.",
    space_after=6
)

add_para(
    "Administrative claims databases and open prescription aggregates are widely used to "
    "study environmental determinants of health. However, when diagnosis-linked records are "
    "unavailable, researchers must select a medication-based outcome that serves as a proxy "
    "for the disease of interest. We quantify how much this choice alone — holding exposure "
    "data, covariates, and model structure constant — changes ecological inference. Using "
    "Japan's National Database (NDB) Open Data and seasonal pollen dispersal as a natural "
    "experiment, we compare three parallel analyses: an allergy-specific outcome "
    "(antihistamines, ophthalmic and nasal preparations), a pharmacologically heterogeneous "
    "comparator (topical anti-inflammatory agents), and a pharmacologically unrelated negative "
    "control (oral hypoglycemic agents). The allergy-specific outcome recovered a robust "
    "pollen signal (R² = 0.469, pollen significant in 7/8 sensitivity specifications), "
    "whereas the heterogeneous comparator attenuated the signal (R² = 0.272, 5/6 "
    "specifications). The negative control showed no pollen association in any specification, "
    "confirming outcome specificity rather than generic confounding.",
    space_after=6
)

add_para(
    "This study makes three contributions relevant to the readership of Pharmacoepidemiology "
    "and Drug Safety: (1) it demonstrates that outcome definition is a primary design "
    "determinant in prescription-based ecological analyses, not a downstream analytic choice; "
    "(2) it applies a novel three-arm comparison framework — disease-aligned, heterogeneous "
    "comparator, and negative control — that can be adopted in future pharmacoepidemiology "
    "research; and (3) it uses Japan's NDB Open Data, a large publicly available national "
    "prescription resource, to illustrate these points at population scale.",
    space_after=6
)

add_para(
    "The analysis plan was preregistered on the Open Science Framework prior to submission "
    "(https://osf.io/yuc4a; registered 18 June 2026). All data used are publicly available "
    "and analysis code is openly available on GitHub (https://github.com/haruki00430/NDB_pollen_allergy) "
    "and archived on Zenodo (https://doi.org/10.5281/zenodo.20747801). This manuscript "
    "follows the STROBE and RECORD reporting guidelines.",
    space_after=6
)

add_para(
    "This work has not been published previously and is not under consideration elsewhere. "
    "A previous version was submitted to the Journal of Clinical Epidemiology "
    "(JCEPI-D-26-00998) but received a desk decision indicating that the topic was "
    "considered more appropriate for a specialist pharmacoepidemiology journal. We believe "
    "Pharmacoepidemiology and Drug Safety is the most appropriate venue for this work.",
    space_after=6
)

add_para(
    "Neither author has any conflicts of interest to declare. This study did not receive "
    "specific external funding. We have completed the required Conflict of Interest disclosure "
    "form and have it available for submission.",
    space_after=6
)

add_para(
    "We hope that this manuscript will be of interest to the readers of "
    "Pharmacoepidemiology and Drug Safety, and we look forward to your consideration.",
    space_after=12
)

add_para("Sincerely,", space_after=6)

add_para("Haruki Saito", bold=True, space_after=0)
add_para("Department of Epidemiology", space_after=0)
add_para("Fukushima Medical University School of Medicine", space_after=0)
add_para("1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan", space_after=0)
add_para("Email: m211039@fmu.ac.jp", space_after=6)

add_para("On behalf of all authors:", space_after=6)
add_para("Haruki Saito and Tetsuya Ohira", space_after=0)

doc.save(DST)
print(f"[OK] カバーレター保存完了: {DST}")
