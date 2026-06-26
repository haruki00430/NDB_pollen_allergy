"""
PDS投稿用 COI (Conflict of Interest) 宣言書 DOCX を作成する。
submission_package_PDS/COI_Statement_PDS.docx として保存。

PDS COI要件:
- PDS指定のICMJE COIフォームを提出する（https://www.icmje.org/disclosure-of-interest/）
- 本スクリプトはICMJEフォームの補助として使用するCOI宣言書を作成する
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DST = (
    r"C:\Users\user\.ag-cursor-common\research_workspace\projects"
    r"\NDB_Research_Hub\projects\NDB_XXX_pollen_allergy_v2\04_Manuscripts"
    r"\submission_package_PDS\COI_Statement_PDS.docx"
)

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(text, bold=False, font_size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    return p

# ── タイトル ─────────────────────────────────────────────────────────────────
add_heading("Conflict of Interest Disclosure Statement", level=1)
add_para(
    "Manuscript title: How Outcome Definition Changes Ecological Pharmacoepidemiological "
    "Inference: A Natural Experiment on Outcome Misclassification",
    font_size=10
)
add_para("Journal: Pharmacoepidemiology and Drug Safety", font_size=10)
add_para("Submission date: 26 June 2026", font_size=10, space_after=12)

# ── 著者 1 ────────────────────────────────────────────────────────────────────
add_heading("Author 1: Haruki Saito", level=2)

add_para("Affiliation:", bold=True, space_after=0)
add_para(
    "Department of Epidemiology, Fukushima Medical University School of Medicine, "
    "1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan",
    space_after=6
)

add_para("Financial support for this work:", bold=True, space_after=0)
add_para("None declared.", space_after=6)

add_para("Grants or research funding:", bold=True, space_after=0)
add_para("None.", space_after=6)

add_para("Honoraria, consulting fees, or other payments:", bold=True, space_after=0)
add_para("None.", space_after=6)

add_para("Employment:", bold=True, space_after=0)
add_para("Fukushima Medical University School of Medicine (graduate student).", space_after=6)

add_para("Stock ownership or options:", bold=True, space_after=0)
add_para("None.", space_after=6)

add_para("Patents:", bold=True, space_after=0)
add_para("None.", space_after=6)

add_para("Other financial relationships:", bold=True, space_after=0)
add_para("None.", space_after=12)

add_para(
    "I certify that the information provided above is complete and accurate.",
    space_after=3
)
add_para("Signature: ______________________________", space_after=0)
add_para("Date: 26 June 2026", space_after=18)

# ── 著者 2 ────────────────────────────────────────────────────────────────────
add_heading("Author 2: Tetsuya Ohira", level=2)

add_para("Affiliation:", bold=True, space_after=0)
add_para(
    "Department of Epidemiology, Fukushima Medical University School of Medicine, "
    "1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan",
    space_after=6
)

add_para("Financial support for this work:", bold=True, space_after=0)
add_para("None declared.", space_after=6)

add_para("Grants or research funding:", bold=True, space_after=0)
add_para("None directly related to this work.", space_after=6)

add_para("Honoraria, consulting fees, or other payments:", bold=True, space_after=0)
add_para("None related to this work.", space_after=6)

add_para("Employment:", bold=True, space_after=0)
add_para(
    "Professor, Department of Epidemiology, Fukushima Medical University School of Medicine.",
    space_after=6
)

add_para("Stock ownership or options:", bold=True, space_after=0)
add_para("None.", space_after=6)

add_para("Patents:", bold=True, space_after=0)
add_para("None.", space_after=6)

add_para("Other financial relationships:", bold=True, space_after=0)
add_para("None.", space_after=12)

add_para(
    "I certify that the information provided above is complete and accurate.",
    space_after=3
)
add_para("Signature: ______________________________", space_after=0)
add_para("Date: 26 June 2026", space_after=18)

# ── 注記 ─────────────────────────────────────────────────────────────────────
add_heading("Notes", level=2)
add_para(
    "This disclosure statement is provided to accompany the ICMJE Conflict of Interest "
    "disclosure forms, which should be obtained from https://www.icmje.org/disclosure-of-interest/ "
    "and completed separately by each author before submission. Each author should complete "
    "and upload an individual ICMJE COI form through the Wiley Research Exchange submission portal.",
    font_size=10
)

doc.save(DST)
print(f"[OK] COI宣言書保存完了: {DST}")
